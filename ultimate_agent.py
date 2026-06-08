import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import requests
from requests.adapters import HTTPAdapter
from urllib3.util import Retry
import json
import operator
import time
import copy # <-- NEW: Needed for HITL evidence editing
import xml.etree.ElementTree as ET
import re
import networkx as nx
from pyvis.network import Network
import streamlit.components.v1 as components
import tempfile
import os
import markdown
import gseapy as gp
import matplotlib.pyplot as plt
from gseapy.plot import gseaplot
from io import BytesIO
from docx import Document
from htmldocx import HtmlToDocx
from typing import TypedDict, List, Dict, Any, Annotated
import py3Dmol
from stmol import showmol
from pydantic import BaseModel, Field
from patsy import dmatrix
# --- NEW RAG IMPORTS ---
from langchain_core.documents import Document as LCDocument # <-- ALIAS FIX
from langchain_community.callbacks import get_openai_callback
from PyPDF2 import PdfReader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS

from openai import OpenAI
from langchain_openai import ChatOpenAI
from langchain_core.messages import SystemMessage, HumanMessage, AIMessage
from langgraph.graph import StateGraph, START, END

from pydeseq2.dds import DeseqDataSet
from pydeseq2.ds import DeseqStats
from stats_engine import run_differential_stats

# ==========================================
# PAGE CONFIGURATION & SECRETS
# ==========================================
st.set_page_config(page_title="Agentic Oncology Orchestrator", layout="wide")

# --- PASSWORD PROTECTION ---
def check_password():
    """Returns `True` if the user had the correct password."""
    def password_entered():
        if st.session_state["password"] == st.secrets["APP_PASSWORD"]:
            st.session_state["password_correct"] = True
            del st.session_state["password"]  # Don't store password
        else:
            st.session_state["password_correct"] = False

    if "password_correct" not in st.session_state:
        st.text_input("🔒 Enter Lab Password", type="password", on_change=password_entered, key="password")
        return False
    elif not st.session_state["password_correct"]:
        st.text_input("🔒 Enter Lab Password", type="password", on_change=password_entered, key="password")
        st.error("😕 Password incorrect")
        return False
    return True

if not check_password():
    st.stop() # Stops the rest of the app from loading until password is correct!

st.title("🧬 Agentic Precision Oncology Pipeline")
st.markdown("Powered by LangGraph, PyDESeq2, OncoKB, and PubMed")

try:
    openai_key = st.secrets["OPENAI_API_KEY"]
    oncokb_key = st.secrets["ONCOKB_API_KEY"]
except KeyError:
    st.error("⚠️ Secrets not found! Please ensure you have a .streamlit/secrets.toml file with your API keys.")
    st.stop()

# --- PHASE 4: RESILIENT HTTP SESSION ---
retry_strategy = Retry(
    total=3,
    backoff_factor=1,
    status_forcelist=[429, 500, 502, 503, 504],
    allowed_methods=["GET", "POST"]
)
adapter = HTTPAdapter(max_retries=retry_strategy)
http_session = requests.Session()
http_session.mount("https://", adapter)
http_session.mount("http://", adapter)

# --- INITIALIZE SESSION STATE (MEMORY) ---
if "run_complete" not in st.session_state:
    st.session_state.run_complete = False
if "messages" not in st.session_state:
    st.session_state.messages = []
if "gathering_complete" not in st.session_state:
    st.session_state.gathering_complete = False
if "agent_state" not in st.session_state:
    st.session_state.agent_state = {}
if "total_tokens" not in st.session_state:
    st.session_state.total_tokens = 0
if "total_cost" not in st.session_state:
    st.session_state.total_cost = 0.0

# --- NEW: GLOBAL STATE INITIALIZATION ---
if "upregulated_df" not in st.session_state:
    st.session_state.upregulated_df = pd.DataFrame()
if "full_results_df" not in st.session_state:
    st.session_state.full_results_df = pd.DataFrame()
if "volcano_fig" not in st.session_state:
    st.session_state.volcano_fig = None
if "ai_targets" not in st.session_state:
    st.session_state.ai_targets = []
if "de_engine" not in st.session_state:
    st.session_state.de_engine = "PyDESeq2"
if "pval_thresh" not in st.session_state:
    st.session_state.pval_thresh = 0.0500
if "log2fc_thresh" not in st.session_state:
    st.session_state.log2fc_thresh = 2.0
if "n_up_pathway" not in st.session_state:
    st.session_state.n_up_pathway = 2
if "n_down_pathway" not in st.session_state:
    st.session_state.n_down_pathway = 1
if "n_outliers" not in st.session_state:
    st.session_state.n_outliers = 1
if "last_cancer_type" not in st.session_state: # <-- NEW
    st.session_state.last_cancer_type = "Melanoma"
if "user_intention" not in st.session_state:
    st.session_state.user_intention = ""
if "cancer_type" not in st.session_state:
    st.session_state.cancer_type = "Melanoma"

# ==========================================
# 1. GRAPH STATE & SCHEMAS
# ==========================================
class AgentState(TypedDict):
    user_prompt: str
    significant_genes: List[Dict[str, Any]]
    plan: List[str]
    gathered_evidence: Annotated[List[Dict[str, Any]], operator.add]
    pathway_data: Dict[str, Any] 
    final_report: str
    custom_knowledge: str 
    analysis_mode: str
    discarded_evidence: List[Dict[str, Any]] 
    ai_filtered_evidence: List[Dict[str, Any]]
    expert_consensus: str # <-- NEW: Holds the multi-agent debate
    is_hitl_run: bool # <-- NEW: Flags if we should pause after gathering
    user_intention: str # <-- NEW: User's explicit research goal for semantic alignment

class Plan(BaseModel):
    steps: List[str] = Field(description="Step-by-step plan of tools to execute.")

# --- NEW: AI SCORER SCHEMA ---
class PaperScore(BaseModel):
    score: int = Field(description="Relevance score from 1 to 10")
    reason: str = Field(description="Short 3-15 word reason (e.g., 'Acronym Collision', 'Strong evidence', 'Wrong Disease')")

# ==========================================
# 2. THE TOOLS (Python Functions)
# ==========================================
@st.cache_data(ttl="1d", show_spinner=False)
def get_gene_info(hugo_symbol):
    """Fetches biological context, gene type, and aliases."""
    url = f"https://mygene.info/v3/query?q=symbol:{hugo_symbol}&fields=name,summary,type_of_gene,alias&species=human"
    try:
        res = http_session.get(url, timeout=10)
        if res.status_code == 200:
            data = res.json()
            if data.get("hits"):
                hit = data["hits"][0]
                # Format aliases nicely whether it's a string or a list
                aliases = hit.get("alias", [])
                if isinstance(aliases, list):
                    aliases = ", ".join(aliases)
                return {
                    "name": hit.get("name", "Unknown"),
                    "type": hit.get("type_of_gene", "Unknown"),
                    "summary": hit.get("summary", "No summary available."),
                    "aliases": aliases
                }
        return {"status": "Gene info not found."}
    except Exception as e:
        return {"status": f"API Error: {str(e)}"}
    
@st.cache_data(ttl="1d", show_spinner=False)
def fetch_normal_tissue_profile(hugo_symbol):
    """Acts as a proxy for the GTEx / Human Protein Atlas databases."""
    llm = ChatOpenAI(model="gpt-5.2", temperature=0, api_key=openai_key)
    
    sys_msg = """You are a Genotype-Tissue Expression (GTEx) and Human Protein Atlas database proxy. 
    Output a strict, 1-sentence summary of where this gene is predominantly expressed in normal, healthy human tissue. 
    Be highly specific (e.g., 'Predominantly expressed in the exocrine pancreas and lactating mammary glands'). 
    If it is ubiquitous across all tissues, explicitly state 'Ubiquitously expressed'."""
    
    try:
        res = llm.invoke([
            SystemMessage(content=sys_msg), 
            HumanMessage(content=f"Gene: {hugo_symbol}")
        ])
        return res.content
    except Exception as e:
        return "GTEx proxy unavailable."

@st.cache_resource(ttl="1d", show_spinner=False)
def run_gsea_analysis(full_df):
    """Runs local GSEA using gseapy on the entire ranked expression profile."""
    # 1. Prepare the ranking metric: -log10(padj) * sign(log2FoldChange)
    df = full_df.dropna(subset=['log2FoldChange', 'padj']).copy()
    df['rank_metric'] = -np.log10(df['padj'] + 1e-300) * np.sign(df['log2FoldChange'])
    
    # 2. Sort from most upregulated to most downregulated
    df = df.sort_values('rank_metric', ascending=False)
    rnk = df[['rank_metric']]
    
    try:
        # 3. Attempt GSEA Prerank locally (FORCED threads=1 to prevent Windows fork bombs!)
        pre_res = gp.prerank(
            rnk=rnk, 
            gene_sets='KEGG_2021_Human',
            threads=1, 
            min_size=5, 
            max_size=1000,
            permutation_num=100, 
            outdir=None, 
            seed=42
        )
        
        res_df = pre_res.res2d
        
        # 4. Filter for significantly enriched pathways (Grab top 10 to ensure we have enough Up and Down options)
        sig_pw = res_df[res_df['FDR q-val'] < 0.05].head(10)
        
        if sig_pw.empty:
            return {"status": "No statistically significant pathways found by GSEA.", "pathways": []}
            
        top_pathways = []
        for idx, row in sig_pw.iterrows():
            # gseapy returns lead genes separated by semicolons
            lead_genes = row['Lead_genes'].split(';')
            top_pathways.append({
                "pathway": row['Term'],
                "p_value": row['NOM p-val'],
                "nes": row.get('NES', 0), # <-- NEW: Track the direction of the pathway!
                "overlapping_genes": lead_genes
            })
            
        return {"status": "Success", "pathways": top_pathways, "gsea_obj": pre_res}
        
    except Exception as e:
        return {"status": f"GSEA failed: {str(e)}"}

@st.cache_data(ttl="1d", show_spinner=False)
def get_onco_data(hugo, alteration, tumor_type):
    url = "https://www.oncokb.org/api/v1/annotate/mutations/byProteinChange"
    params = {"hugoSymbol": hugo, "alteration": alteration, "tumorType": tumor_type}
    headers = {"accept": "application/json", "Authorization": f"Bearer {oncokb_key}"}
    
    try:
        response = http_session.get(url, params=params, headers=headers, timeout=10)
        if response.status_code == 200:
            data = response.json()
            treatments = data.get('treatments', [])
            if not treatments: return {"status": "No drug entries found."}
            
            results = []
            for treatment in treatments:
                drugs = [d.get('drugName', '') for d in treatment.get('drugs', [])]
                results.append({
                    "drugName": ", ".join(drugs), 
                    "levelOfEvidence": treatment.get('level', 'Unknown'),
                    "pmids": treatment.get('pmids', [])
                })
            return {"status": "Success", "drugs": results}
        return {"status": f"OncoKB Error: {response.status_code}"}
    except Exception as e:
        return {"status": f"Request failed: {str(e)}"}

@st.cache_data(ttl="1d", show_spinner=False)
def search_pubmed(gene, tumor_type, mode="Clinical Triage", aliases="", interactors=None, user_intention=""):
    if interactors is None: interactors = []
    
    alias_query = ""
    if aliases and aliases != "Unknown":
        # NEW: Only keep aliases longer than 3 characters to prevent massive acronym pollution!
        alias_list = [a.strip() for a in aliases.split(',') if len(a.strip()) > 3][:2]
        if alias_list:
            alias_query = " OR " + " OR ".join([f"{a}[TIAB]" for a in alias_list])

    # --- NEW: UNCONDITIONAL NETWORK PULL FOR DISCOVERY ---
    if "Discovery" in mode and interactors:
        network_nodes = [gene] + interactors
        network_query_str = " OR ".join([f"{n}[TIAB]" for n in network_nodes])
        broad_query = f"({network_query_str}{alias_query}) AND {tumor_type}[TIAB]"
        prov_step_1 = f"**Phase 1 (Broad Network Pull):** Expanded query to include STRING interactors: `[{broad_query}]`."
    else:
        broad_query = f"({gene}[TIAB]{alias_query}) AND {tumor_type}[TIAB]"
        prov_step_1 = f"**Phase 1 (Broad Target Pull):** PubMed query `[{broad_query}]`."
    
    search_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi"
    search_params = {"db": "pubmed", "term": broad_query, "retmode": "json", "retmax": 40}
    
    try:
        res = http_session.get(search_url, params=search_params, timeout=10)
        if res.status_code != 200: return {"status": f"Search Error: {res.status_code}"}
            
        id_list = res.json().get("esearchresult", {}).get("idlist", [])
        if not id_list: return {"status": "No experimental literature found."}
            
        time.sleep(0.5) 
        fetch_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/efetch.fcgi"
        fetch_params = {"db": "pubmed", "id": ",".join(id_list), "retmode": "xml"}
        
        fetch_res = http_session.get(fetch_url, params=fetch_params, timeout=15)
        if fetch_res.status_code != 200: return {"status": "PubMed Fetch Error"}
            
        papers = []
        root = ET.fromstring(fetch_res.content)
        for article in root.findall('.//PubmedArticle'):
            pmid = article.find('.//PMID').text if article.find('.//PMID') is not None else "Unknown"
            title = article.find('.//ArticleTitle').text if article.find('.//ArticleTitle') is not None else "No Title"
            
            abstract_nodes = article.findall('.//AbstractText')
            abstract_text = " ".join([node.text for node in abstract_nodes if node.text]) if abstract_nodes else ""
                
            if abstract_text: 
                papers.append({"PMID": pmid, "Title": title, "Abstract": abstract_text[:1500]})
                
        if not papers: return {"status": "No abstracts available to embed."}

        # --- 2. SEMANTIC FILTER (FAISS VECTOR DB) ---
        st.markdown(f"      -> Embedding {len(papers)} abstracts into FAISS for {gene} network...")
        docs = [LCDocument(page_content=f"Title: {p['Title']}\nAbstract: {p['Abstract']}", metadata=p) for p in papers]
        
        embeddings = OpenAIEmbeddings(api_key=openai_key)
        vectorstore = FAISS.from_documents(docs, embeddings)
        
        if "Discovery" in mode:
            semantic_query = f"Novel biomarkers, signaling pathways, lipid metabolism, immunotherapy targets, and resistance mechanisms in {tumor_type}."
        else:
            semantic_query = f"FDA approved targeted therapy, survival outcomes, and clinical trial results for {tumor_type}."
            
        # --- NEW: INTENT-DRIVEN RAG ALIGNMENT ---
        if user_intention.strip():
            semantic_query += f" Explicit focus requested on: {user_intention.strip()}."
            
        retriever = vectorstore.as_retriever(search_kwargs={"k": 10})
        relevant_docs = retriever.invoke(semantic_query)
        top_papers = [{"PMID": d.metadata["PMID"], "Title": d.metadata["Title"], "Abstract": d.metadata["Abstract"]} for d in relevant_docs]
        
        # --- 3. PROVENANCE LOGGING ---
        provenance = [
            prov_step_1 + f" Yielded {len(id_list)} candidates.",
            f"**Phase 2 (Semantic Sorting):** Embedded {len(papers)} valid abstracts into FAISS Vector DB.",
            f"**Phase 3 (Concept Retrieval):** Extracted top 10 papers by mathematical proximity to: *'{semantic_query}'*.",
            f"**Phase 4 (Expert Review):** Top semantic matches passed to the AI Scorer for strict clinical triage."
        ]
        
        return {"status": "Success", "papers": top_papers, "provenance": provenance}
        
    except Exception as e:
        return {"status": f"Request failed: {str(e)}"}

@st.cache_data(ttl="1d", show_spinner=False)
def search_clinical_trials(gene, tumor_type):
    url = "https://clinicaltrials.gov/api/v2/studies"
    query = f"{gene} AND {tumor_type}"
    params = {"query.cond": query, "filter.overallStatus": "RECRUITING", "pageSize": 3}
    
    try:
        res = http_session.get(url, params=params, timeout=10)
        if res.status_code == 200:
            data = res.json()
            studies = data.get("studies", [])
            if not studies:
                return {"status": "No recruiting trials found."}
                
            trials = []
            for study in studies:
                protocol = study.get("protocolSection", {})
                ident = protocol.get("identificationModule", {})
                design = protocol.get("designModule", {}) 
                
                nct_id = ident.get("nctId", "Unknown NCT")
                title = ident.get("briefTitle", "No Title")
                phase = ", ".join(design.get("phases", ["Phase Unknown"])) 
                
                trials.append({"NCT_ID": nct_id, "Title": title, "Phase": phase})
                
            time.sleep(0.5)
            return {"status": "Success", "trials": trials}
            
        return {"status": f"ClinicalTrials Error: {res.status_code}"}
    except Exception as e:
        return {"status": f"Request failed: {str(e)}"}

@st.cache_data(ttl="1d", show_spinner=False)
def get_protein_interactions(hugo_symbol):
    """Fetches top 3 interacting proteins from STRING DB (Guilt by Association)."""
    # 9606 is the NCBI taxonomy ID for Homo sapiens
    url = f"https://string-db.org/api/json/network?identifiers={hugo_symbol}&species=9606&limit=3"
    try:
        res = http_session.get(url, timeout=10)
        if res.status_code == 200:
            data = res.json()
            if not data:
                return {"status": "No interactions found."}
            
            interactors = []
            for edge in data:
                # Get the protein that is NOT our query gene
                neighbor = edge.get("preferredName_B") if edge.get("preferredName_A") == hugo_symbol else edge.get("preferredName_A")
                if neighbor and neighbor not in interactors:
                    interactors.append(neighbor)
            
            # Keep only the top 3 unique neighbors
            interactors = interactors[:3]
            return {"status": "Success", "interacting_proteins": interactors}
        return {"status": f"STRING API Error: {res.status_code}"}
    except Exception as e:
        return {"status": f"Request failed: {str(e)}"}

@st.cache_data(ttl="1d", show_spinner=False)
def fetch_target_tractability(hugo_symbol):
    """Fetches Druggability (Tractability) and Essentiality from the Open Targets API."""
    url = "https://api.platform.opentargets.org/api/v4/graphql"
    
    # NEW: 'entity' changed to 'object' to match the V4 schema
    query = """
    query targetSearch($queryString: String!) {
      search(queryString: $queryString, entityNames: ["target"]) {
        hits {
          object {
            ... on Target {
              id
              approvedSymbol
              tractability {
                label
                modality
                value
              }
              depMapEssentiality {
                screens {
                  depmapId
                  diseaseFromSource
                }
              }
            }
          }
        }
      }
    }
    """
    variables = {"queryString": hugo_symbol}
    try:
        res = http_session.post(url, json={"query": query, "variables": variables}, timeout=10)
        if res.status_code == 200:
            data = res.json()
            hits = data.get("data", {}).get("search", {}).get("hits", [])
            for hit in hits:
                # NEW: Python now extracts from the 'object' dictionary
                obj = hit.get("object", {})
                if obj and obj.get("approvedSymbol") == hugo_symbol:
                    
                    # 1. Parse Tractability
                    tractability = obj.get("tractability") or []
                    is_druggable = False
                    modalities = []
                    for t in tractability:
                        if t.get("value") == True:
                            is_druggable = True
                            modalities.append(f"{t.get('modality')} ({t.get('label')})")
                    
                    # 2. Parse Essentiality (Handling the OpenTargets List Schema)
                    essentiality_data = obj.get("depMapEssentiality") or []
                    is_essential = False
                    essential_screens = 0
                    
                    # OpenTargets returns a list, so we safely check if it has items and grab the first one
                    if isinstance(essentiality_data, list) and len(essentiality_data) > 0:
                        screens = essentiality_data[0].get("screens", [])
                        essential_screens = len(screens)
                        is_essential = essential_screens > 0
                    
                    # --- NEW: EXPLICIT 'NO DATA' DECLARATION ---
                    if not is_druggable and not is_essential:
                        status_msg = "Target exists in OpenTargets, but contains ZERO tractability or DepMap essentiality data."
                    else:
                        status_msg = "Success"

                    return {
                        "status": status_msg,
                        "is_druggable": is_druggable,
                        "tractability_buckets": modalities[:5],
                        "is_depmap_essential": is_essential,
                        "essential_cell_lines": essential_screens
                    }
            return {"status": f"Target '{hugo_symbol}' not found in OpenTargets Database. Verify HGNC symbol."}
        return {"status": f"API Error: {res.status_code}"}
    except Exception as e:
        return {"status": f"Request failed: {str(e)}"}

def process_pdf_for_rag(pdf_file):
    """Reads a PDF, splits it into chunks, and builds a FAISS vector database."""
    reader = PdfReader(pdf_file)
    raw_text = ""
    for page in reader.pages:
        extracted = page.extract_text()
        if extracted:
            raw_text += extracted
            
    # CRITICAL: Prevent the database from crashing if the PDF is just images!
    if not raw_text.strip():
        return None 
        
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=150,
        length_function=len
    )
    chunks = text_splitter.split_text(raw_text)
    
    embeddings = OpenAIEmbeddings(api_key=openai_key)
    vectorstore = FAISS.from_texts(chunks, embeddings)
    
    return vectorstore

@st.cache_data(ttl="1d", show_spinner=False)
def get_uniprot_id(hugo_symbol):
    """Maps a HGNC Gene Symbol to a UniProt ID using MyGene.info"""
    url = f"https://mygene.info/v3/query?q=symbol:{hugo_symbol}&fields=uniprot&species=human"
    try:
        res = http_session.get(url, timeout=10)
        if res.status_code == 200:
            data = res.json()
            if data.get("hits"):
                hit = data["hits"][0]
                uniprot = hit.get("uniprot", {})
                if "Swiss-Prot" in uniprot:
                    return uniprot["Swiss-Prot"]
                elif "TrEMBL" in uniprot:
                    return uniprot["TrEMBL"][0] if isinstance(uniprot["TrEMBL"], list) else uniprot["TrEMBL"]
        return None
    except Exception:
        return None

@st.cache_data(ttl="1d", show_spinner=False)
def fetch_alphafold_structure(uniprot_id):
    """Fetches the 3D coordinates dynamically from the AlphaFold EBI API"""
    api_url = f"https://alphafold.ebi.ac.uk/api/prediction/{uniprot_id}"
    try:
        api_res = http_session.get(api_url, timeout=10)
        if api_res.status_code == 200:
            data = api_res.json()
            if data and isinstance(data, list):
                file_url = data[0].get("pdbUrl")
                file_format = "pdb"
                if not file_url:
                    file_url = data[0].get("cifUrl")
                    file_format = "cif"
                if file_url:
                    struct_res = http_session.get(file_url, timeout=15)
                    if struct_res.status_code == 200:
                        return struct_res.text, file_format
        return None, None
    except Exception:
        return None, None

@st.cache_data(show_spinner=False)
def get_uniprot_binding_sites(uniprot_id):
    """Autonomously fetches known Active Sites and Ligand Binding Pockets from UniProt"""
    url = f"https://rest.uniprot.org/uniprotkb/{uniprot_id}.json"
    try:
        res = http_session.get(url, timeout=10)
        if res.status_code == 200:
            features = res.json().get("features", [])
            target_residues = []
            for f in features:
                if f.get("type") in ["Binding site", "Active site"]:
                    loc = f.get("location", {})
                    start = loc.get("start", {}).get("value")
                    end = loc.get("end", {}).get("value")
                    if start and end: target_residues.extend(list(range(start, end + 1)))
                    elif start: target_residues.append(start)
            return [str(r) for r in set(target_residues)]
        return []
    except: return []

def extract_residue_number(mutation_string):
    """Extracts the numeric position from a string like 'V600E'"""
    if not mutation_string: return None
    match = re.search(r'\d+', mutation_string)
    return match.group() if match else None

def render_mutated_protein(structure_data, file_format="pdb", highlight_residues=None):
    """Renders protein. Highlights pockets if provided, else falls back to Confidence coloring."""
    view = py3Dmol.view(width=800, height=500)
    view.addModel(structure_data, file_format)
    
    if highlight_residues:
        view.setStyle({'model': -1}, {"cartoon": {'color': 'lightgrey'}})
        if isinstance(highlight_residues, str) or isinstance(highlight_residues, int):
            highlight_residues = [str(highlight_residues)]
        view.addStyle({'resi': highlight_residues}, {'sphere': {'color': 'red', 'radius': 1.2}})
        view.addStyle({'resi': highlight_residues}, {'stick': {'colorscheme': 'blueCarbon'}})
    else:
        view.setStyle({'model': -1}, {"cartoon": {'colorscheme': {'prop':'b','gradient': 'roygb','min':50,'max':90}}})
        
    view.zoomTo()
    return view

@st.cache_data(show_spinner=False)
def fetch_visual_network(hugo_symbol, max_nodes=15):
    """Fetches a larger interacting network specifically for the UI Graph"""
    url = f"https://string-db.org/api/json/network?identifiers={hugo_symbol}&species=9606&limit={max_nodes}"
    try:
        res = http_session.get(url, timeout=10)
        if res.status_code == 200:
            return res.json()
        return []
    except Exception:
        return []

def build_pyvis_graph(central_gene, edges_data):
    """Builds a physics-based interactive graph with anti-overlap repulsion"""
    G = nx.Graph()
    for edge in edges_data:
        node_a = edge.get("preferredName_A")
        node_b = edge.get("preferredName_B")
        score = edge.get("score", 0)
        if node_a and node_b and score > 0.4:
            G.add_edge(node_a, node_b, weight=score)

    net = Network(height="600px", width="100%", bgcolor="#0E1117", font_color="white")
    
    for node in G.nodes():
        if node == central_gene:
            net.add_node(node, label=node, color="#EF553B", size=30, shape="star")
        else:
            size = 15 + (G.degree(node) * 2)
            net.add_node(node, label=node, color="#636EFA", size=size)

    for edge in G.edges(data=True):
        net.add_edge(edge[0], edge[1], value=edge[2]['weight'], color="#4A4A4A")

    # --- THE PHYSICS UPGRADE ---
    # We force 'repulsion' to keep nodes far apart so labels never overlap
    net.repulsion(node_distance=200, central_gravity=0.1, spring_length=200, spring_strength=0.05, damping=0.09)
    return net

# --- BENCH-TO-CLOUD WET LAB SCHEMAS ---
class sgRNA(BaseModel):
    target_exon: str = Field(description="Which exon to target (e.g., Exon 2)")
    sequence: str = Field(description="20bp RNA sequence")
    pam: str = Field(description="PAM sequence (e.g., NGG)")
    off_target_risk: str = Field(description="Low, Medium, or High")

class PrimerPair(BaseModel):
    target: str = Field(description="What this primer amplifies")
    forward: str = Field(description="Forward primer sequence (18-22bp)")
    reverse: str = Field(description="Reverse primer sequence (18-22bp)")
    tm: str = Field(description="Melting temperature (e.g., 60 C)")
    amplicon_size: int = Field(description="Expected amplicon size in bp")

class WetLabManifest(BaseModel):
    rationale: str = Field(description="1-sentence explanation of why we are knocking out this gene.")
    sgrnas: List[sgRNA] = Field(description="Top 2 sgRNA designs for CRISPR knockout.")
    primers: List[PrimerPair] = Field(description="qPCR primers to validate the knockout.")

@st.cache_data(show_spinner=False)
def design_validation_experiment(gene_symbol, cancer_type):
    """Autonomously designs CRISPR sgRNAs and qPCR primers."""
    llm = ChatOpenAI(model="gpt-5.2", temperature=0.1, api_key=openai_key)
    structured_llm = llm.with_structured_output(WetLabManifest)
    
    sys_msg = """You are an expert Molecular Biologist and CRISPR designer.
    Design a realistic, immediate wet-lab validation experiment to knock out this gene using CRISPR-Cas9.
    CRITICAL RULES:
    1. Output strictly in the requested JSON format.
    2. sgRNA sequences MUST be exactly 20 nucleotides + a valid PAM (e.g., CGG, TGG).
    3. Primers MUST be 18-22 nucleotides with ~50-60% GC content.
    4. Always include a housekeeping gene primer pair (e.g., ACTB or GAPDH) as a control."""
    
    prompt = f"Target Gene: {gene_symbol}\nCancer Context: {cancer_type}\nDesign the CRISPR and qPCR validation manifest."
    try:
        response = structured_llm.invoke([SystemMessage(content=sys_msg), HumanMessage(content=prompt)])
        return response.model_dump() # Return as dict for Streamlit caching!
    except Exception as e:
        return None

# ==========================================
# 3. LANGGRAPH NODES
# ==========================================
def planner_node(state: AgentState):
    st.write("📝 **[Node: Planner]** Architecting 8-gate clinical data gathering strategy...")
    llm = ChatOpenAI(model="gpt-5.2", temperature=0, api_key=openai_key)
    structured_llm = llm.with_structured_output(Plan)
    
    sys_msg = """You are an expert Clinical Bioinformatics Planner. 
    Analyze the user prompt and genes. Output a step-by-step plan to gather data.
    Available Tools: 
    1. 'Biological Context Gate' (Maps baseline disease-to-gene variants and handles initial disease context onboarding)
    2. 'GTEx Tissue Gate' (Audits healthy tissue expression baselines to identify potential off-tumor toxicities)
    3. 'OpenTargets Gate' (Evaluates drug tractability modalities and CRISPR cellular essentiality knockout scores)
    4. 'OncoKB Gate' (Cross-references FDA targeted therapies, actionability tiers, and resistance profiles)
    5. 'STRING Network Gate' (Constructs and clusters functional protein-protein interaction networks for pathway context)
    6. 'PubMed & AI Triage Gate' (Performs intent-driven semantic literature reviews to uncover hidden molecular mechanisms)
    7. 'UniProt Structural Gate' (Audits 3D protein structures and subcellular localization to verify cell-surface vs. intracellular expression)
    8. 'Clinical Trials Gate' (Identifies actively recruiting human clinical trials and patient cohort criteria)
    
    Directive: The execution plan must be highly comprehensive, structurally robust, and clinical-grade. Ensure every gate relevant to solving the user's specific biological or therapeutic intent is actively leveraged. Do not omit necessary analytical steps; prioritize absolute thoroughness and translational utility for the clinical oncology team."""
    
    context = f"User Prompt: {state.get('user_prompt')}\nGenes: {state.get('significant_genes')}"
    response = structured_llm.invoke([SystemMessage(content=sys_msg), HumanMessage(content=context)])
    
    return {"plan": response.steps}

def executor_node(state: AgentState):
    st.write("⚙️ **[Node: Executor]** Engaging multi-omic tool execution...")
    plan_text = " ".join(state.get("plan", [])).lower()
    genes = state.get("significant_genes", [])
    new_evidence = []
    
    # COPY the discarded evidence list to prevent mutating the graph state dictionary directly!
    ai_filtered_evidence_list = state.get("ai_filtered_evidence", []).copy()
    
    for gene_info in genes:
        hugo = gene_info.get("hugo")
        alt = gene_info.get("alteration")
        tumor_type = gene_info.get("tumor_type")
        source_tag = gene_info.get("source", "Unknown Source")
        
        report = {"gene": hugo, "alteration": alt, "source": source_tag, "biology": {}, "evidence": {}}
        
        # --- 1. BIOLOGICAL CONTEXT GATE ---
        try:
            st.write(f"   - 🧬 **[Biological Context Gate]** Fetching baseline definition for {hugo}...")
            gene_context = get_gene_info(hugo)
        except Exception as e:
            gene_context = {"name": "Unknown", "type": "Unknown", "summary": "API Unavailable.", "aliases": ""}
            
        report["biology"] = gene_context
        
        # --- 2. GTEX TISSUE GATE ---
        try:
            st.write(f"   - 🔬 **[GTEx Tissue Gate]** Profiling normal tissue distribution for {hugo}...")
            tissue_profile = fetch_normal_tissue_profile(hugo)
            report["biology"]["normal_tissue_gtex"] = tissue_profile
        except Exception as e:
            report["biology"]["normal_tissue_gtex"] = "GTEx proxy unavailable."
            
        # --- 3. UNIPROT STRUCTURAL GATE ---
        try:
            st.write(f"   - 🧩 **[UniProt Structural Gate]** Auditing 3D sub-cellular localization for {hugo}...")
            uniprot_id = get_uniprot_id(hugo)
            if uniprot_id:
                pockets = get_uniprot_binding_sites(uniprot_id)
                report["evidence"]["UniProt_Pockets"] = {
                    "has_defined_pockets": len(pockets) > 0, 
                    "residue_count": len(pockets),
                    "note": "If > 0, this protein has known druggable active sites/pockets."
                }
        except Exception as e:
            report["evidence"]["UniProt_Pockets"] = {"status": f"Database Connection Unavailable: {str(e)}"}
            
        # --- 4. OPEN TARGETS GATE ---
        try:
            st.write(f"   - 🎯 **[OpenTargets Gate]** Fetching Tractability & Essentiality for {hugo}...")
            report["evidence"]["OpenTargets"] = fetch_target_tractability(hugo)
        except Exception as e:
            report["evidence"]["OpenTargets"] = {"status": f"Database Connection Unavailable: {str(e)}"}
        
        # --- 5. ONCOKB GATE ---
        if "oncokb" in plan_text:
            try:
                st.write(f"   - 💊 **[OncoKB Gate]** Cross-referencing FDA approvals for {hugo}...")
                report["evidence"]["OncoKB"] = get_onco_data(hugo, alt, tumor_type)
            except Exception as e:
                report["evidence"]["OncoKB"] = {"status": f"Database Connection Unavailable: {str(e)}"}
            
        # --- 6. STRING NETWORK GATE ---
        if "Discovery" in state.get("analysis_mode", "Clinical Triage"):
            try:
                st.write(f"   - 🕸️ **[STRING Network Gate]** Constructing interaction network for {hugo}...")
                report["evidence"]["STRING_Interactions"] = get_protein_interactions(hugo)
            except Exception as e:
                report["evidence"]["STRING_Interactions"] = {"status": f"Database Connection Unavailable: {str(e)}"}
            
        # --- 7. PUBMED & AI TRIAGE GATE ---
        if "pubmed" in plan_text:
            try:
                st.write(f"   - 📚 **[PubMed & AI Triage Gate]** Performing intent-driven semantic literature review for {hugo}...")
                # EXTRACT THE NETWORK WE JUST FETCHED (safe defaults if network fetching failed)
                interactors = report.get("evidence", {}).get("STRING_Interactions", {}).get("interacting_proteins", [])
                
                pubmed_data = search_pubmed(
                    hugo, 
                    tumor_type, 
                    mode=state.get("analysis_mode", "Clinical Triage"),
                    aliases=gene_context.get("aliases", ""),
                    interactors=interactors, # <-- Pass the network into the tool!
                    user_intention=state.get("user_intention", "") # <-- NEW: Pass intent to semantic search
                )
                
                # Save the Semantic Search Provenance Log!
                report["evidence"]["PubMed_Provenance"] = pubmed_data.get("provenance", [])
                
                # --- AI RELEVANCE SCORER (OVERSAMPLE & FILTER) ---
                if pubmed_data.get("status") == "Success" and pubmed_data.get("papers"):
                    st.write(f"      -> 🤖 Grading literature relevance for {hugo}...")
                    grader_llm = ChatOpenAI(model="gpt-5.2", temperature=0, api_key=openai_key).with_structured_output(PaperScore)
                    
                    candidate_papers = pubmed_data["papers"]
                    good_papers = []
                    
                    bio_name = gene_context.get('name', 'Unknown')
                    bio_summary = gene_context.get('summary', 'No summary available.')
                    
                    for p in candidate_papers:
                        if len(good_papers) >= 3:
                            break # We found 3 good papers! Stop grading to save OpenAI tokens.
                            
                        # Tell the scorer to accept network interactions!
                        network_str = f" OR its immediate functional network ({', '.join(interactors)})" if interactors else ""
                        
                        eval_prompt = f"""
                        Evaluate this abstract's relevance to the target {hugo} ({bio_name}){network_str} in {tumor_type}.
                        Biological Function of {hugo}: {bio_summary}
                        
                        CRITICAL RUBRIC:
                        - Score 1-4: Acronym collision (e.g., the gene symbol refers to a drug/procedure), completely unrelated disease, or irrelevant biology.
                        - Score 5-10: Relevant. The primary gene {hugo} {network_str} is mentioned in a functional, prognostic, or therapeutic context. (Score highly if it validates the target's network!).
                        
                        Title: {p['Title']}
                        Abstract: {p['Abstract'][:800]}
                        """
                        
                        try:
                            score_result = grader_llm.invoke([
                                SystemMessage(content="You are an expert oncology peer-reviewer. Output strict JSON grading the paper's relevance."),
                                HumanMessage(content=eval_prompt)
                            ])
                            p["AI_Score"] = score_result.score
                            p["AI_Reason"] = score_result.reason
                            
                            if score_result.score >= 5:
                                good_papers.append(p)
                            else:
                                # Toss it into the AI's trash can!
                                ai_filtered_evidence_list.append({
                                    "Gene": hugo,
                                    "Score": score_result.score,
                                    "Reason": score_result.reason,
                                    "Title": p["Title"],
                                    "PMID": p["PMID"]
                                })
                                
                        except Exception as e:
                            print(f"Scoring failed for {hugo}: {e}")
                            p["AI_Score"] = "?"
                            p["AI_Reason"] = "Error"
                            good_papers.append(p) # Keep it if scoring fails just to be safe
                    
                    # Replace the massive list of 10 papers with ONLY the good ones
                    pubmed_data["papers"] = good_papers
                            
                report["evidence"]["PubMed"] = pubmed_data
            except Exception as e:
                report["evidence"]["PubMed"] = {"status": f"Database Connection Unavailable: {str(e)}"}
                report["evidence"]["PubMed_Provenance"] = ["PubMed API or Semantic Embedding Pipeline Failed."]

        # --- 8. CLINICAL TRIALS GATE ---
        if "clinicaltrials" in plan_text or "trials" in plan_text:
            try:
                st.write(f"   - 🏥 **[Clinical Trials Gate]** Fetching actively recruiting trials for {hugo}...")
                report["evidence"]["ClinicalTrials"] = search_clinical_trials(hugo, tumor_type)
            except Exception as e:
                report["evidence"]["ClinicalTrials"] = {"status": f"Database Connection Unavailable: {str(e)}"}
            
        new_evidence.append(report)
        time.sleep(1.5) # <-- Prevents local script flooding, network throttling handled by adapter
        
    return {"gathered_evidence": new_evidence, "ai_filtered_evidence": ai_filtered_evidence_list}

def clinical_review_node(state: AgentState):
    st.write("🧑‍⚕️ **[Node: Clinical Review]** Cross-referencing findings with multi-agent clinical tumor board...")
    llm = ChatOpenAI(model="gpt-5.2", temperature=0.3, api_key=openai_key)
    
    prompt = f"""
    You are hosting a clinical tumor board. Review this preliminary data for {state.get('user_prompt')}.
    Clean Evidence: {json.dumps(state.get('gathered_evidence'))}
    Pathways: {json.dumps(state.get('pathway_data'))}
    
    First, speak as a MOLECULAR PATHOLOGIST: In 1 paragraph, evaluate the tissue context and biological plausibility. CRITICAL SANITY CHECK: You MUST explicitly compare the gene's "normal_tissue_gtex" profile (found in the Clean Evidence) against the user's specified Cancer Type. If the gene is a canonical marker for a completely different tissue type (e.g., a pancreas gene in a melanoma sample), you MUST flag this as a probable lineage artifact or sample contamination. Furthermore, if the user prompt mentions "CAR-T", "CAR", "antibody", "ADC", or "surface", explicitly audit the OpenTargets tractability and biological summary to determine sub-cellular localization. State clearly whether the target is a transmembrane/surface-bound receptor or an intracellular/nuclear factor.
    Second, speak as a MEDICAL ONCOLOGIST: In 1 paragraph, evaluate the druggability and clinical trial viability. You MUST explicitly reference the OpenTargets Tractability and DepMap Essentiality data provided in the Clean Evidence to justify your assessment. If the Pathologist flags a tissue mismatch, advise extreme caution regarding clinical utility. Additionally, if the Pathologist determines the target is intracellular but the user proposed a surface-modality (like CAR-T or ADC), you MUST issue a strict warning invalidating the target for that specific modality.
    Third, speak as a BIOINFORMATICS AUDITOR: In 1 paragraph, audit the PubMed literature for Acronym Collisions. (e.g., if the gene symbol is 'CEL' but the abstract is talking about 'CEL cells' or 'Celastrol', or 'PPL' meaning a polymer). If you detect a collision or biologically disconnected paper, explicitly name it and command the Medical Writer to ignore it.
    """
    response = llm.invoke([HumanMessage(content=prompt)])
    return {"expert_consensus": response.content}

def writer_node(state: AgentState):
    st.write("✍️ **[Node: Writer]** Synthesizing evidence into clinical-grade report...")
    llm = ChatOpenAI(model="gpt-5.2", temperature=0.2, api_key=openai_key)
    
    analysis_mode = state.get("analysis_mode", "Clinical Triage")
    
    if "Discovery" in analysis_mode:
        sys_msg = """You are an expert Systems Biologist and Bioinformatics AI.
        Write a beautiful, pathway-centric scientific report answering the user's prompt. 
        
        CRITICAL GUARDRAILS:
        1. TONE AND STYLE: Write confidently as if authoring a published review article in a high-impact oncology journal. NEVER break the fourth wall (do not say "The user provided" or "The AI found"). Use bullet points occasionally to break up dense mechanistic explanations. Use **bold text** for gene names.
        2. BIOLOGICAL TRIAGE: Explicitly dismiss pseudogenes and ncRNAs as non-coding artifacts.
        3. ACRONYM COLLISIONS & LITERATURE LIMITATIONS: Be highly aware of literature false-positives. If PubMed returns acronym collisions or papers that only discuss network neighbors (and not the primary target), weave this critique naturally into your academic prose (e.g., "While some literature attributes X to Y, this largely reflects confounding nomenclature..." or "Current evidence is restricted to proximal network nodes rather than direct target validation..."). Do NOT use robotic sub-headers like "Sanity Check".
        4. SYSTEMS APPROACH: Discuss genes conceptually within their pathways.
        5. GUILT BY ASSOCIATION: If a target lacks direct literature or tractability, evaluate its "STRING_Interactions". Propose backdoor therapeutic strategies using its direct neighbors.
        6. TISSUE CONTEXT: If the Expert Consensus flags lineage mismatches (e.g., breast genes in melanoma, or stromal admixture), state this clearly once in the executive summary, and frame the rest of the report around resolving this spatial uncertainty.
        7. SUB-CELLULAR LOCALIZATION & SAFETY PROFILE: For every target evaluated, you MUST actively audit the UniProt structural data for sub-cellular localization and the GTEx data for healthy tissue expression. If an immunotherapy modality like CAR-T or an antibody drug conjugate (ADC) is requested, a target CANNOT be classified as Tier 1 unless it is explicitly verified as extracellular/surface-bound and has a low healthy-tissue cross-reactivity profile.
        
        REQUIRED REPORT STRUCTURE:
        ## 📊 Executive Summary
        [Write a concise 3-4 sentence high-level overview. YOU MUST START by explicitly explaining to the reader WHY these specific genes were selected (e.g., "These targets were autonomously selected because they are the primary drivers of the [Insert Top KEGG Pathway Here] cluster identified in the patient's differential expression profile.") Then summarize actionable next steps.]
        
        ## 🕸️ Systems Biology & Pathway Dysregulation
        [Write a multi-paragraph synthesis of the KEGG pathway data. How do these networks (and their overlapping genes) interact to drive the tumor microenvironment, metabolic reprogramming, or immune evasion?]
        
        ## 🔬 Targetable Hubs & Translational Risk Tiers
        [Synthesize the literature conceptually, but you MUST categorize each evaluated gene into one of the following Translational Risk Tiers based strictly on its OpenTargets Tractability and DepMap Essentiality data:]
        
        * **🟢 Tier 1: Actionable Hubs (Low Risk)**: The gene is classified as highly Druggable, highly Essential, AND explicitly verified via UniProt structural data as localized to the extracellular plasma membrane (essential for cell-surface modalities like CAR-T/ADCs) with clean GTEx healthy-tissue baselines. Ready for translational validation.
        * **🟡 Tier 2: Network Dependencies (Moderate Risk)**: The gene is biologically relevant but lacks direct OpenTargets druggability (e.g., intracellular and non-essential). However, its STRING interactors or pathways are actionable. Intervene at the network level.
        * **🟠 Tier 3: Orphan Signals (High Risk)**: The gene is Not Tractable, Not Essential in DepMap, lacks literature, and has no actionable interactors, OR targets proposed for cell-surface therapies that UniProt data flags as strictly intracellular, nuclear, or mitochondrial (representing a structural mismatch). Highly speculative; requires orthogonal wet-lab validation.
        * **🔴 Tier 4: Probable Artifacts (Do Not Pursue)**: Pseudogenes, lineage mismatches (e.g., breast genes in melanoma), or acronym collisions.
        
        [Discuss the genes under their appropriate Tier headers.]
        
        ## 🏥 Translational Outlook
        [Summarize any relevant trials, or state that these novel network targets currently lack specific recruiting trials. Explicitly comment on WHY this biological connection is novel and HOW it is biologically plausible based on the Pathologist/Oncologist consensus.]
        
        ### 🧪 Recommended Next Experimental Steps
        [Provide 3-4 bullet points on how a wet-lab researcher should experimentally validate these findings. CRITICAL: Your final bullet point MUST inform the user that they can generate a ready-to-order CRISPR/qPCR Wet-Lab Manifest for any of these targets using the "Bench-to-Cloud Validation Designer" at the bottom of the dashboard.]
        """
    else:
        sys_msg = """You are an expert Clinical Oncology Medical Writer.
        Write a beautiful, multi-paragraph scientific report answering the user's prompt.
        
        CRITICAL CLINICAL TRIAGE RULES:
        1. Standard of Care: Level 1 or 2 evidence.
        2. Repurposing: Level 3 or 4 evidence.
        3. Dismiss pseudogenes/ncRNAs using the biology context.
        
        REQUIRED REPORT STRUCTURE:
        First, provide a brief summary of the overarching biological pathways driving this tumor:
        ### 🕸️ Pathway & Network Dysregulation
        - [Write a 2-3 sentence summary based on the KEGG pathway_data provided. Mention the overlapping genes.]
        
        Then, for EVERY gene sequentially, use this exact structure:
        ## [Gene Name] ([Alteration])
        
        ### 🎯 Target Tractability & Essentiality (OpenTargets)
        [CRITICAL: If the OpenTargets status indicates 'ZERO tractability' or 'not found', you MUST explicitly write: "OpenTargets was queried but currently holds no tractability or essentiality data for this target."]
        - **Druggability:** [Summarize modality buckets, or state none]
        - **Essentiality:** [State if it is essential in DepMap screens, or state none]
        - **Structural Vulnerability:** [CRITICAL: Check the 'UniProt_Pockets' data. Explicitly state if the target has defined druggable active sites/pockets, or if it lacks them.]
        - **Sub-cellular Localization (UniProt/GTEx):** [Explicitly state if the protein is localized on the cell surface/transmembrane vs intracellular, and note if GTEx flags any critical healthy-tissue expression liabilities.]

        ### 💊 OncoKB Therapeutics
        - **Standard of Care (On-Label):** [Drug Name] (PMIDs: [List])
        - **Repurposing Opportunities (Off-Label):** [Drug Name] (PMIDs: [List])
        
        ### 🔬 Experimental Literature
        - **[Study Topic]:** [Summary] (PMID: [Number])
        
        ### 🏥 Actively Recruiting Trials
        - **[[NCT ID]](https://clinicaltrials.gov/study/[NCT ID]):** [Phase] - [Trial Title]
        
        Do not deviate from this structure."""
    
    user_context = f"User Prompt: {state.get('user_prompt')}\nPathway Data: {json.dumps(state.get('pathway_data', {}))}\nExpert Consensus: {state.get('expert_consensus')}\nGathered Evidence: {json.dumps(state.get('gathered_evidence'))}\nCustom Lab Protocols: {state.get('custom_knowledge', 'None provided.')}"
    
    response = llm.invoke([
        SystemMessage(content=sys_msg),
        HumanMessage(content=user_context)
    ])
    
    st.write("✅ **Final report successfully written.**")
    return {"final_report": response.content}

def route_start(state: AgentState):
    """Routes to Phase 2 if evidence already exists, else starts Phase 1."""
    # FIX: Rely on 'plan' existence AND the 'is_hitl_run' un-pause flag to know we are in Phase 2
    # This prevents infinite loops if the user threw ALL papers in the trash!
    if state.get("plan") and len(state.get("plan")) > 0 and not state.get("is_hitl_run", True):
        return "clinical_review"
    return "planner"

def route_executor(state: AgentState):
    """Pauses the graph if HITL is enabled; otherwise continues to clinical review."""
    if state.get("is_hitl_run", False):
        return END
    return "clinical_review"

workflow = StateGraph(AgentState)
workflow.add_node("planner", planner_node)
workflow.add_node("executor", executor_node)
workflow.add_node("clinical_review", clinical_review_node) # <-- NEW
workflow.add_node("writer", writer_node)

# --- NEW: CONDITIONAL ROUTING EDGES ---
workflow.add_conditional_edges(START, route_start)
workflow.add_edge("planner", "executor")
workflow.add_conditional_edges("executor", route_executor)
workflow.add_edge("clinical_review", "writer")
workflow.add_edge("writer", END)

orchestrator = workflow.compile()

# ==========================================
# STREAMLIT FRONTEND & UI (VERSION 2.0)
# ==========================================

col1, col2 = st.columns([1, 2])

with col1:
    st.subheader("1. Session Data")
    counts_file = st.file_uploader("Upload RNA Counts (CSV)", type=["csv"])
    metadata_file = st.file_uploader("Upload Metadata (CSV)", type=["csv"])
    
    st.markdown("---")
    st.subheader("Optional: DNA Mutational Profile")
    dna_file = st.file_uploader("Upload DNA Variants (CSV with 'Gene' and 'Alteration' columns)", type=["csv"])
    
    # --- NEW: Dynamic Covariate Selection ---
    condition_col = "condition" # Fallbacks
    batch_col = "None"
    
    if metadata_file is not None:
        # Peek at the metadata columns without locking up memory
        temp_meta = pd.read_csv(metadata_file, nrows=0) 
        meta_cols = temp_meta.columns.tolist()
        
        # CRITICAL FIX: Rewind the file pointer back to the beginning!
        metadata_file.seek(0)
        
        st.markdown("---")
        st.subheader("2. Experimental Design")
        col_a, col_b = st.columns(2)
        with col_a:
            condition_col = st.selectbox("Primary Contrast (e.g., Tumor vs Normal)", meta_cols, index=meta_cols.index("condition") if "condition" in meta_cols else 0)
        with col_b:
            batch_col = st.selectbox("Batch Covariate (Optional)", ["None"] + meta_cols)
            
    st.markdown("---")
    st.subheader("3. Configuration & Target Roster")
    
    # --- Unified Configuration Form ---
    with st.form(key="target_settings_form"):
        st.markdown("#### 📝 Research Context")
        user_intention_input = st.text_area(
            "Research Goal / Intention (Optional)", 
            value=st.session_state.user_intention,
            placeholder="E.g., 'I am specifically looking for novel lipid metabolism targets...' or 'Focus on resistance mechanisms to BRAF inhibitors.'",
            help="Guide the AI's literature search and synthesis. Leave blank for standard clinical triage."
        )
        cancer_type_input = st.text_input("Cancer Type (e.g., Melanoma, NSCLC)", value=st.session_state.cancer_type)
        
        st.markdown("#### 📈 Statistical Cutoffs")
        de_engine_input = st.selectbox("Differential Expression Engine", ["PyDESeq2", "RPKM/T-Test"], index=0 if st.session_state.de_engine == "PyDESeq2" else 1)
        pval_thresh_input = st.number_input("P-Value Cutoff", min_value=0.0001, max_value=0.1000, value=st.session_state.pval_thresh, step=0.0100, format="%.4f")
        log2fc_thresh_input = st.slider("Log2FC Threshold (Absolute)", min_value=0.0, max_value=10.0, value=st.session_state.log2fc_thresh, step=0.5)
        
        st.markdown("#### 🎯 Target Selection Roster")
        st.write("Configure how many targets the AI should extract from the different biological layers:")
        col_r1, col_r2, col_r3 = st.columns(3)
        
        with col_r1:
            n_up_input = st.number_input("Upregulated Pathway Drivers", min_value=0, max_value=5, value=st.session_state.n_up_pathway, help="Genes driving the dominant GSEA networks.")
        with col_r2:
            n_down_input = st.number_input("Downregulated Biomarkers", min_value=0, max_value=5, value=st.session_state.n_down_pathway, help="Suppressed genes from inverse GSEA networks (Synthetic Lethality).")
        with col_r3:
            n_out_input = st.number_input("Lone Wolves (Outliers)", min_value=0, max_value=5, value=st.session_state.n_outliers, help="Highest statistical spikes, regardless of pathway.")
        
        submitted = st.form_submit_button("Apply Configuration & Generate Plot")
        
    if submitted:
        # --- PHASE 4 FIX: STALE CONTEXT RESET ---
        # Moving this inside the form submission guarantees we only purge state when actively applying a new context!
        if st.session_state.last_cancer_type != cancer_type_input:
            st.session_state.upregulated_df = pd.DataFrame()
            st.session_state.full_results_df = pd.DataFrame()
            st.session_state.ai_targets = []
            st.session_state.gathering_complete = False
            st.session_state.run_complete = False
            st.session_state.agent_state = {}
            st.session_state.last_cancer_type = cancer_type_input
            
        # Explicitly bind mutated values to state ONLY when form is submitted to prevent race conditions
        st.session_state.user_intention = user_intention_input
        st.session_state.cancer_type = cancer_type_input
        st.session_state.de_engine = de_engine_input
        st.session_state.pval_thresh = pval_thresh_input
        st.session_state.log2fc_thresh = log2fc_thresh_input
        st.session_state.n_up_pathway = n_up_input
        st.session_state.n_down_pathway = n_down_input
        st.session_state.n_outliers = n_out_input
        
    # Extract from state for safe downstream usage (retained during unrelated UI rerenders)
    user_intention = st.session_state.user_intention
    cancer_type = st.session_state.cancer_type
    de_engine = st.session_state.de_engine
    pval_thresh = st.session_state.pval_thresh
    log2fc_thresh = st.session_state.log2fc_thresh
    n_up_pathway = st.session_state.n_up_pathway
    n_down_pathway = st.session_state.n_down_pathway
    n_outliers = st.session_state.n_outliers
    
    top_n_genes = n_up_pathway + n_down_pathway + n_outliers # Keep total count for downstream logic
    update_plot_btn = submitted # Map the submit button to the downstream mathematical trigger

    st.markdown("---")
    st.subheader("4. Clinical Context & AI Triage")
    
    analysis_mode = st.radio("Analysis Mode", ["Clinical Triage (Known Targets)", "Biomarker Discovery (Novel Targets)"])
    
    st.markdown("### 🧑‍⚕️ Clinical Safety & Evidence")
    hitl_toggle = st.toggle("⏸️ Enable Human-in-the-Loop (Review evidence before report generation)", value=True)
    baseline_toggle = st.toggle("⚖️ Head-to-Head Baseline (Compare Agent vs. Vanilla LLM)", value=False)
    
    # NEW: Dynamic button text!
    if hitl_toggle:
        btn_text = "⏸️ Step 1: Gather Evidence for Review"
    else:
        btn_text = "🚀 Run Full AI Clinical Triage"
        
    # --- PHASE 2: UI GUARD LOGIC ---
    # Check if PyDESeq2/T-Test data has been successfully generated
    data_ready = not st.session_state.upregulated_df.empty
    
    if not data_ready:
        st.info("⚠️ Please upload your data and click 'Apply Configuration & Generate Plot' first to populate actionable targets.")
        
    run_button = st.button(
        btn_text, 
        use_container_width=True, 
        type="primary",
        disabled=not data_ready # Physically lock the button to prevent out-of-order execution!
    )
    
    # --- NEW RAG UI ---
    st.markdown("---")
    st.subheader("5. Custom Knowledge (Optional)")
    uploaded_pdf = st.file_uploader("Upload Lab Protocols/Guidelines (PDF)", type=["pdf"])

with col2:
    st.subheader("Interactive Volcano Plot")
    
    # Only run the heavy math if files are uploaded AND the update button was clicked
    if counts_file and metadata_file and update_plot_btn:
        counts_df = pd.read_csv(counts_file, index_col=0)
        metadata_df = pd.read_csv(metadata_file, index_col=0)
        
        with st.spinner(f"Calculating Differential Expression using {de_engine}..."):
            # Determine the design formula strings based on user selection
            if batch_col != "None" and batch_col != condition_col:
                design_factors = [batch_col, condition_col]
                edge_formula = f"~{batch_col} + {condition_col}"
            else:
                design_factors = condition_col
                edge_formula = f"~{condition_col}"

            # Auto-detect the contrast levels from the primary column
            unique_levels = metadata_df[condition_col].dropna().unique()
            level_1 = unique_levels[0]
            level_2 = unique_levels[1] if len(unique_levels) > 1 else unique_levels[0]
            
            mock_active = st.session_state.get("use_mock_mode", False)

            # --- 1. DATA VALIDATION & ALIGNMENT GATEKEEPER ---
            is_transposed = False
            # Standard RNA-seq data often has genes as rows and samples as columns.
            # PyDESeq2 strictly expects samples as rows and genes as columns.
            if set(metadata_df.index).issubset(set(counts_df.columns)):
                counts_df = counts_df.T
                is_transposed = True
                
            # Assert exact matching after potential transposition
            try:
                assert set(metadata_df.index).issubset(set(counts_df.index)), "Sample IDs in metadata do not match counts matrix."
            except AssertionError as e:
                st.error(f"🚨 Data Alignment Error: {str(e)}. Please check your CSV files.")
                st.stop()
                
            # Align and subset counts to match metadata order perfectly
            counts_df = counts_df.loc[metadata_df.index]
            
            # --- 2. DATA TYPE SAFETY ---
            # PyDESeq2 strictly requires integer matrices. Handle NaNs and floats safely.
            counts_df = counts_df.fillna(0).round().astype(int)
            
            # Remove zero-variance genes to prevent math convergence failures
            counts_df = counts_df.loc[:, counts_df.var(axis=0) > 0]

            # --- THE OVERRIDE ---
            if mock_active:
                fallback_counts = counts_df.T if is_transposed else counts_df
                output = run_differential_stats(fallback_counts, metadata_df, condition_col, level_1, level_2, mock_mode=True)
                results_df = output["results_df"]
            else:
                if de_engine == "PyDESeq2":
                    try:
                        # --- 1. DEFENSIVE NA & TYPE CLEANING ---
                        # Fill any missing values with 0 and explicitly force integer types (mandatory for PyDESeq2)
                        counts_df = counts_df.fillna(0).astype(int)

                        # --- 2. DYNAMIC SAMPLE-PERCENTAGE FILTERING ---
                        # Define thresholds: A gene must have at least 10 reads in at least 20% of the total samples
                        min_reads_per_sample = 10
                        sample_percentage_threshold = 0.20

                        # Calculate the minimum number of samples required based on current dataset size (at least 1)
                        min_required_samples = max(1, int(len(counts_df) * sample_percentage_threshold))

                        # Identify columns (genes) where the count is >= 10 in at least the minimum required samples
                        genes_passing_filter = (counts_df >= min_reads_per_sample).sum(axis=0) >= min_required_samples

                        # Filter the matrix columns to keep only the passing genes
                        counts_df = counts_df.loc[:, genes_passing_filter]

                        # --- 3. INITIALIZE MULTI-THREADED ENGINE ---
                        # Ensure n_cpus=4 is active to parallelize across our system cores
                        dds = DeseqDataSet(
                            counts=counts_df,
                            metadata=metadata_df,
                            design=edge_formula,
                            refit_cooks=True,
                            n_cpus=4
                        )
                        dds.deseq2()
                        stat_res = DeseqStats(dds, contrast=[condition_col, level_1, level_2], n_cpus=4)
                        stat_res.summary()
                        results_df = stat_res.results_df
                    except Exception as e:
                        # --- 3. DEFENSIVE MATH EXECUTION (FALLBACK) ---
                        st.warning(f"⚠️ PyDESeq2 Math Convergence Error: {str(e)}. Automatically falling back to robust T-Test Engine...")
                        fallback_counts = counts_df.T if is_transposed else counts_df
                        output = run_differential_stats(
                            counts_df=fallback_counts,
                            metadata_df=metadata_df,
                            condition_col=condition_col,
                            test_cond=level_1,
                            ctrl_cond=level_2,
                            mock_mode=False
                        )
                        results_df = output["results_df"]
                        
                elif de_engine == "RPKM/T-Test":
                    # --- PRODUCTION RPKM T-TEST PIPELINE ---
                    fallback_counts = counts_df.T if is_transposed else counts_df
                    output = run_differential_stats(
                        counts_df=fallback_counts,
                        metadata_df=metadata_df,
                        condition_col=condition_col,
                        test_cond=level_1,
                        ctrl_cond=level_2,
                        mock_mode=False
                    )
                    results_df = output["results_df"]
            
        plot_df = results_df.dropna(subset=['padj', 'log2FoldChange']).copy()
        plot_df['-log10(padj)'] = -np.log10(plot_df['padj'] + 1e-300)
        
        conditions = [
            (plot_df['padj'] < pval_thresh) & (plot_df['log2FoldChange'] > log2fc_thresh), 
            (plot_df['padj'] < pval_thresh) & (plot_df['log2FoldChange'] < -log2fc_thresh)
        ]
        plot_df['Significance'] = np.select(conditions, ['Upregulated', 'Downregulated'], default='Not Significant')
        
        # Save all upregulated genes to memory for the Actionability Filter
        st.session_state.upregulated_df = plot_df[plot_df['Significance'] == 'Upregulated'].sort_values(by='padj')
        
        # NEW: Save the FULL results dataframe for the GSEA math engine
        st.session_state.full_results_df = plot_df.copy()

        # Generate a clean map of the tumor
        plot_title = "Gene Expression Volcano Plot"
        if "Discovery" in analysis_mode:
            plot_title += "<br><sup>⭐ Targets selected via Pathway-Cluster analysis (bypassing isolated statistical spikes)</sup>"

        fig = px.scatter(
            plot_df, x='log2FoldChange', y='-log10(padj)', color='Significance', 
            color_discrete_map={
                'Upregulated': '#EF553B', 
                'Downregulated': '#636EFA', 'Not Significant': '#4A4A4A' 
            },
            hover_name=plot_df.index,
            render_mode='webgl',
            title=plot_title  # <-- NEW: Explicitly explaining the stars!
        )

        # --- NEW: Changed lines to white ---
        fig.add_hline(y=-np.log10(pval_thresh), line_dash="dash", line_color="white")
        fig.add_vline(x=log2fc_thresh, line_dash="dash", line_color="white")
        fig.add_vline(x=-log2fc_thresh, line_dash="dash", line_color="white")
        fig.update_layout(height=500)
        
        st.session_state.volcano_fig = fig # Save plot to memory
        
        # --- PHASE 3 FIX: FORCE UI SYNC ---
        # Instantly flush the state and redraw the UI so the AI execution button unlocks immediately
        st.rerun()
        
    # Always display the plot if it exists in memory, even if they clicked a different button!
    if st.session_state.volcano_fig:
        st.plotly_chart(st.session_state.volcano_fig, use_container_width=True)
        
        if len(st.session_state.ai_targets) > 0:
            formatted_genes = ", ".join([f"`{gene}`" for gene in st.session_state.ai_targets])
            st.success(f"✅ **{len(st.session_state.ai_targets)} Targets identified:** {formatted_genes}")
        else:
            st.warning("⚠️ **No targets selected.** Adjust your statistical cutoffs and update the plot.")
    elif not counts_file or not metadata_file:
        st.info("👈 Upload data and click 'Apply Configuration & Generate Plot' to begin.")

# ==========================================
# EXECUTE THE AI GRAPH
# ==========================================
if run_button and counts_file and metadata_file:
    st.markdown("---")
    st.subheader("🤖 AI Clinical Report")
    
    # --- NEW: CLUSTER-FIRST TARGET SELECTION ---
    ACTIONABLE_GENES = ["BRAF", "EGFR", "KRAS", "PIK3CA", "ERBB2", "ALK", "ROS1", "MET", "RET", "NTRK1", "NTRK2", "NTRK3", "BRCA1", "BRCA2", "KIT", "PDGFRA", "FGFR1", "FGFR2", "FGFR3", "IDH1", "IDH2", "CDK4", "CDK6", "PTEN", "MTOR", "CTNNB1", "TP53"]
    
    up_df = st.session_state.get("upregulated_df", pd.DataFrame())
    if up_df.empty:
        st.error("⚠️ No upregulated genes found. Please lower your P-Value or Log2FC thresholds in the Volcano Plot first.")
        st.stop()
        
    with st.spinner("🧠 Recruiting Hybrid Target Roster via Local GSEA..."):
        full_df = st.session_state.get("full_results_df", pd.DataFrame())
        if full_df.empty:
            st.error("⚠️ Full results missing. Please re-run the Volcano plot.")
            st.stop()
            
        # Prepare the pools to pull from
        if "Discovery" in analysis_mode:
            # Filter out known actionables AND noisy Ribosomal/Mitochondrial housekeeping genes
            gsea_input_df = full_df[
                (~full_df.index.isin(ACTIONABLE_GENES)) & 
                (~full_df.index.str.match(r'^(RPL|RPS|MT-)')) # <-- NEW: The Ribosome Scrubber
            ]
        else:
            gsea_input_df = full_df
            
        up_df_pool = gsea_input_df[gsea_input_df['log2FoldChange'] > 0].sort_values(by='padj')
        down_df_pool = gsea_input_df[gsea_input_df['log2FoldChange'] < 0].sort_values(by='padj')
        extreme_df_pool = gsea_input_df.sort_values(by='padj') # Absolute highest significance

        # Run GSEA
        pathway_results = run_gsea_analysis(gsea_input_df)
        
        # --- THE FIX: Extract the complex math object to Streamlit memory, then delete it from the AI payload ---
        st.session_state.gsea_obj = pathway_results.pop("gsea_obj", None)
        
        cluster_targets = []
        roster_metadata = [] # Keeps track of WHY the AI picked them
        
        up_pathways = [pw for pw in pathway_results.get("pathways", []) if pw.get("nes", 0) > 0]
        down_pathways = [pw for pw in pathway_results.get("pathways", []) if pw.get("nes", 0) < 0]

        # --- 1. UPREGULATED DRIVERS ---
        up_count = 0
        for pw in up_pathways:
            for g in pw["overlapping_genes"]:
                if up_count >= n_up_pathway: break
                if g in up_df_pool.index and g not in cluster_targets:
                    cluster_targets.append(g)
                    roster_metadata.append({"gene": g, "source": f"Upregulated Driver ({pw['pathway']})", "alteration": "Overexpressed"})
                    up_count += 1
                    
        # Pad Up Drivers if GSEA found too few
        for g in up_df_pool.index:
            if up_count >= n_up_pathway: break
            if g not in cluster_targets:
                cluster_targets.append(g)
                roster_metadata.append({"gene": g, "source": "Upregulated Outlier (Padding)", "alteration": "Overexpressed"})
                up_count += 1

        # --- 2. DOWNREGULATED BIOMARKERS ---
        down_count = 0
        for pw in down_pathways:
            for g in pw["overlapping_genes"]:
                if down_count >= n_down_pathway: break
                if g in down_df_pool.index and g not in cluster_targets:
                    cluster_targets.append(g)
                    roster_metadata.append({"gene": g, "source": f"Downregulated Biomarker ({pw['pathway']})", "alteration": "Loss of Expression"})
                    down_count += 1
                    
        # Pad Down Biomarkers if GSEA found too few
        for g in down_df_pool.index:
            if down_count >= n_down_pathway: break
            if g not in cluster_targets:
                cluster_targets.append(g)
                roster_metadata.append({"gene": g, "source": "Downregulated Outlier (Padding)", "alteration": "Loss of Expression"})
                down_count += 1

        # --- 3. LONE WOLVES (OUTLIERS) ---
        outlier_count = 0
        for g in extreme_df_pool.index:
            if outlier_count >= n_outliers: break
            if g not in cluster_targets:
                cluster_targets.append(g)
                direction = "Overexpressed" if full_df.loc[g, 'log2FoldChange'] > 0 else "Loss of Expression"
                roster_metadata.append({"gene": g, "source": "Lone Wolf (Statistical Outlier)", "alteration": direction})
                outlier_count += 1

        # Save to memory
        st.session_state.ai_targets = cluster_targets
        st.session_state.roster_metadata = roster_metadata
        
        st.success(f"🧬 **Hybrid Target Roster Locked:** {', '.join(st.session_state.ai_targets)}")

        # --- NEW: HIGHLIGHT TARGETS ON THE VOLCANO PLOT ---
        full_df = st.session_state.get("full_results_df", pd.DataFrame())
        if st.session_state.volcano_fig is not None and not full_df.empty:
            # Safely grab only the targets that actually exist in the full dataframe
            valid_targets = [g for g in st.session_state.ai_targets if g in full_df.index]
            target_df = full_df.loc[valid_targets]
            
            for idx, row in target_df.iterrows():
                # Dynamic color: Red for Upregulated, Blue for Downregulated
                bg_color = "#EF553B" if row['log2FoldChange'] > 0 else "#636EFA"
                
                st.session_state.volcano_fig.add_annotation(
                    x=row['log2FoldChange'],
                    y=row['-log10(padj)'],
                    text=f"⭐ {idx}",
                    showarrow=True,
                    arrowhead=2,
                    arrowsize=1,
                    arrowwidth=2,
                    arrowcolor="white",
                    font=dict(color="white", size=12, weight="bold"),
                    bgcolor=bg_color,
                    bordercolor="white",
                    borderwidth=1
                )
    
    # --- NEW: RAG PDF PROCESSING (BULLETPROOF VERSION) ---
    rag_context = ""
    if uploaded_pdf is not None:
        try:
            with st.spinner("📚 Reading uploaded Lab Protocol into Vector Database..."):
                vectorstore = process_pdf_for_rag(uploaded_pdf)
                
                if vectorstore is None:
                    st.warning("⚠️ Could not read text from this PDF (it might be a scanned image). Proceeding without custom knowledge.")
                else:
                    retriever = vectorstore.as_retriever(search_kwargs={"k": 3})
                    query = f"Protocols, guidelines, and context for {cancer_type} or genes: {', '.join(st.session_state.ai_targets)}"
                    docs = retriever.invoke(query)
                    rag_context = "\n\n".join([d.page_content for d in docs])
                    st.success("✅ Custom Knowledge Base loaded and queried!")
                    
        except Exception as e:
            st.warning(f"⚠️ PDF Database Error: {str(e)}. Proceeding using only public data.")
    
    with st.status("🧠 Live Agent Thought Trace (Glass Box)", expanded=True) as status:
        structured_genes = []
        dna_gene_names = []
        
        # 1. Parse Optional DNA Mutations (Highest Priority for FDA Drugs)
        if dna_file is not None:
            try:
                dna_df = pd.read_csv(dna_file)
                
                # NEW: Hard stop if the columns are wrong
                if 'Gene' not in dna_df.columns or 'Alteration' not in dna_df.columns:
                    st.error("🚨 CRITICAL ERROR: Your DNA CSV must contain exactly two columns named 'Gene' and 'Alteration'. Please fix your file and re-upload.")
                    st.stop() # This instantly halts the app to protect clinical safety!
                    
                for _, row in dna_df.iterrows():
                    gene_name = str(row['Gene']).strip()
                    dna_gene_names.append(gene_name)
                    structured_genes.append({
                        "hugo": gene_name,
                        "alteration": str(row['Alteration']).strip(),
                        "tumor_type": cancer_type,
                        "source": "DNA Mutation (Level 1/2 Priority)"
                    })
                dna_file.seek(0)
            except Exception as e:
                st.error(f"🚨 CRITICAL ERROR: Could not read the DNA file: {str(e)}")
                st.stop()
        
        # 2. Add RNA Hybrid Roster Targets
        for target in st.session_state.roster_metadata:
            structured_genes.append({
                "hugo": target["gene"],
                "alteration": target["alteration"], 
                "tumor_type": cancer_type,
                "source": target["source"] # This tells the AI if it's a Pathway Driver, Biomarker, or Lone Wolf!
            })
            
        # 3. Smart Prompt Generation (Handling both DNA and RNA)
        if "Discovery" in analysis_mode:
            base_task = f"Analyze the following dysregulated genes ({', '.join(st.session_state.ai_targets)}) in {cancer_type} as potential novel biomarkers or immunotherapeutic targets. Pay close attention to their directionality (Overexpressed vs. Loss of Expression)."
            if dna_gene_names:
                base_task += f" Also contextualize the presence of these specific DNA mutations: {', '.join(dna_gene_names)}."
        else:
            base_task = f"Find established targeted therapies for {cancer_type} patients."
            if dna_gene_names:
                base_task += f" CRITICAL: Prioritize finding OncoKB Level 1/2 FDA-approved therapies for the following DNA mutations: {', '.join(dna_gene_names)}."
            if st.session_state.ai_targets:
                base_task += f" Secondary: Evaluate the following dysregulated RNA targets: {', '.join(st.session_state.ai_targets)}. Note their directionality in your analysis."

        # Safely inject the User's Free-Text Intention without overriding the core task
        if user_intention.strip():
            prompt_text = f"USER'S SPECIFIC RESEARCH INTENTION: '{user_intention.strip()}'\n\nCORE SYSTEM TASK: {base_task}"
        else:
            prompt_text = base_task

        initial_state = {
            "user_prompt": prompt_text,
            "significant_genes": structured_genes,
            "plan": [],
            "gathered_evidence": [],
            "pathway_data": pathway_results, # <-- NEW: Pass the pre-calculated cluster data!
            "final_report": "",
            "custom_knowledge": rag_context, 
            "analysis_mode": analysis_mode,
            "discarded_evidence": [], 
            "ai_filtered_evidence": [],
            "expert_consensus": "",
            "is_hitl_run": hitl_toggle, # <-- NEW: Inject HITL context for graph router!
            "user_intention": st.session_state.get("user_intention", "") # <-- NEW: Inject intent for FAISS search
        }
        
        # Save settings for the Vanilla Baseline execution
        st.session_state.run_baseline = baseline_toggle
        st.session_state.base_cancer_type = cancer_type
        st.session_state.base_prompt = prompt_text
        
        # --- PHASE 1: GATHERING (The Executor) ---
        with st.status("🧬 Executing Precision Oncology Graph...", expanded=True) as status:
            with get_openai_callback() as cb:
                if hitl_toggle:
                    # 🤖 TRUE AGENTIC INVOCATION (Phase 1 Only)
                    final_output_state = orchestrator.invoke(initial_state)
                else:
                    # 🤖 TRUE AGENTIC INVOCATION (Freight Train Full Pipeline)
                    final_output_state = orchestrator.invoke(initial_state)
                    
                # Securely unpack the returned final graph state back into Streamlit memory
                st.session_state.agent_state = final_output_state
                
                # Accumulate the costs
                st.session_state.total_tokens += cb.total_tokens
                st.session_state.total_cost += cb.total_cost
            status.update(label="✅ Graph Execution Paused/Complete", state="complete", expanded=False)
            
        st.session_state.gathering_complete = True
        st.session_state.run_complete = False # Reset in case of a re-run
        
        if not hitl_toggle:
            # --- NEW: VANILLA BASELINE EXECUTION ---
            if st.session_state.get("run_baseline"):
                with st.status("⚖️ Generating Vanilla LLM Baseline...", expanded=True):
                    st.markdown("Bypassing OpenTargets, OncoKB, and RAG...")
                    vanilla_llm = ChatOpenAI(model="gpt-5.2", temperature=0.2, api_key=openai_key)
                    v_sys = "You are a clinical oncology assistant. Write a report using only your training data. Do not use tools."
                    v_prompt = f"Task: {st.session_state.base_prompt}\nTarget Genes: {', '.join(st.session_state.ai_targets)}\nCancer: {st.session_state.base_cancer_type}"
                    try:
                        v_res = vanilla_llm.invoke([SystemMessage(content=v_sys), HumanMessage(content=v_prompt)])
                        st.session_state.baseline_report = v_res.content
                        st.markdown("✅ Vanilla Baseline complete.")
                    except Exception as e:
                        st.session_state.baseline_report = f"Vanilla LLM Error: {e}"

            st.session_state.run_complete = True
            # FIX: Use safe .get() unpacking to prevent KeyError if generation fails
            st.session_state.final_report = st.session_state.agent_state.get("final_report", "⚠️ Report generation failed.")
            st.session_state.plan = st.session_state.agent_state.get("plan", [])
            st.session_state.pathway_data = st.session_state.agent_state.get("pathway_data", {})
            
        st.rerun() # NEW: Forces Streamlit to cleanly switch to the Pause menu!

# --- PHASE 1.5: THE HUMAN-IN-THE-LOOP PAUSE ---
if st.session_state.get("gathering_complete") and not st.session_state.get("run_complete") and hitl_toggle:
    st.markdown("---")
    st.subheader("⏸️ Human-in-the-Loop: Review Evidence")
    st.info("The AI has gathered the following PubMed literature. Uncheck any irrelevant papers before generating the final clinical report.")
    
    # Flatten the nested PubMed papers into a simple list for the dataframe
    flat_papers = []
    for g_idx, g_data in enumerate(st.session_state.agent_state.get("gathered_evidence", [])):
        papers = g_data.get("evidence", {}).get("PubMed", {}).get("papers", [])
        for p_idx, p in enumerate(papers):
            flat_papers.append({
                "Keep": True,
                "Score (1-10)": p.get("AI_Score", "?"),     # <-- NEW
                "AI Reason": p.get("AI_Reason", "N/A"),     # <-- NEW
                "Gene": g_data["gene"],
                "Title": p["Title"],
                "PMID": p["PMID"],
                "_g_idx": g_idx,  
                "_p_idx": p_idx   
            })
            
    if flat_papers:
        df_papers = pd.DataFrame(flat_papers)
        # Render the interactive Data Editor!
        edited_df = st.data_editor(
            df_papers[["Keep", "Score (1-10)", "AI Reason", "Gene", "Title", "PMID"]], 
            hide_index=True, 
            use_container_width=True,
            disabled=["Score (1-10)", "AI Reason", "Gene", "PMID", "Title"] 
        )
    else:
        st.info("💡 **Literature Triage:** The AI reviewed the retrieved literature but determined none of the papers established a direct, functional link between these specific targets and the selected disease. This may represent a highly novel biological connection, OR it may indicate that these genes are not functionally relevant to this specific cancer lineage (e.g., a data mismatch).")
        edited_df = pd.DataFrame()
    # NEW: Show what the AI automatically discarded
    ai_discarded = st.session_state.agent_state.get("ai_filtered_evidence", [])
    if ai_discarded:
        with st.expander("🤖 AI Pre-Filtered Literature (Auto-Discarded)"):
            st.info("The AI evaluated up to 10 papers per gene. The following papers scored < 5 and were automatically excluded.")
            for doc in ai_discarded:
                st.markdown(f"- **{doc['Gene']}** (Score: {doc['Score']}): *{doc['Title']}* - Reason: `{doc['Reason']}`")
        
    # --- THE FINAL TRIGGER ---
    if st.button("🚀 Step 2: Approve Evidence & Synthesize Report", type="primary", use_container_width=True):
        with st.markdown("✍️ **[NODE: Writer]** Synthesizing the final clinical report..."):
            approved_evidence = copy.deepcopy(st.session_state.agent_state["gathered_evidence"])
            discarded_papers = [] # <-- NEW: Temporary list for trash
            
            if not edited_df.empty:
                # Clear out the original papers
                for g_data in approved_evidence:
                    if "PubMed" in g_data.get("evidence", {}) and "papers" in g_data["evidence"]["PubMed"]:
                        g_data["evidence"]["PubMed"]["papers"] = []
                
                # Loop through the table to sort checked vs unchecked
                for i, row in edited_df.iterrows():
                    # RE-ADDED: Safely unpack the pointers to the original data payload
                    g_idx = flat_papers[i]["_g_idx"]
                    p_idx = flat_papers[i]["_p_idx"]
                    original_paper = st.session_state.agent_state["gathered_evidence"][g_idx]["evidence"]["PubMed"]["papers"][p_idx]
                    
                    if row["Keep"]:
                        # Keep it for the report
                        approved_evidence[g_idx]["evidence"]["PubMed"]["papers"].append(original_paper)
                    else:
                        # Toss it in the trash can
                        discarded_papers.append({
                            "Gene": flat_papers[i]["Gene"],
                            "Title": original_paper.get("Title", "Unknown Title"),
                            "PMID": original_paper.get("PMID", "Unknown PMID")
                        })
                            
            # Save the clean evidence AND the trash back to the AI's brain
            current_state = st.session_state.agent_state
            current_state["gathered_evidence"] = approved_evidence
            current_state["discarded_evidence"] = discarded_papers
            current_state["is_hitl_run"] = False # <-- NEW: Un-pause for Phase 2 synthesis!
            
            # --- PHASE 2: TUMOR BOARD & WRITING ---
            with st.status("🧬 Resuming Precision Oncology Graph Synthesis...", expanded=True) as status:
                with get_openai_callback() as cb:
                    # 🤖 TRUE AGENTIC INVOCATION (Phase 2 Only)
                    final_output_state = orchestrator.invoke(current_state)
                    
                    # Unpack the state cleanly back into Streamlit memory
                    st.session_state.agent_state = final_output_state
                    
                # Accumulate the costs
                st.session_state.total_tokens += cb.total_tokens
                st.session_state.total_cost += cb.total_cost
                status.update(label="✅ Graph Synthesis Complete", state="complete", expanded=False)
        
        # --- NEW: VANILLA BASELINE EXECUTION (HITL) ---
        if st.session_state.get("run_baseline"):
            with st.status("⚖️ Generating Vanilla LLM Baseline...", expanded=True):
                st.markdown("Bypassing OpenTargets, OncoKB, and RAG...")
                vanilla_llm = ChatOpenAI(model="gpt-5.2", temperature=0.2, api_key=openai_key)
                v_sys = "You are a clinical oncology assistant. Write a report using only your training data. Do not use tools."
                v_prompt = f"Task: {st.session_state.base_prompt}\nTarget Genes: {', '.join(st.session_state.ai_targets)}\nCancer: {st.session_state.base_cancer_type}"
                try:
                    v_res = vanilla_llm.invoke([SystemMessage(content=v_sys), HumanMessage(content=v_prompt)])
                    st.session_state.baseline_report = v_res.content
                    st.markdown("✅ Vanilla Baseline complete.")
                except Exception as e:
                    st.session_state.baseline_report = f"Vanilla LLM Error: {e}"

        # Mark as finished and refresh the page to show the results
        st.session_state.run_complete = True
        # FIX: Use safe .get() unpacking
        st.session_state.final_report = st.session_state.agent_state.get("final_report", "⚠️ Report generation failed.")
        st.session_state.plan = st.session_state.agent_state.get("plan", [])
        st.session_state.pathway_data = st.session_state.agent_state.get("pathway_data", {})
        st.rerun()
        
# ==========================================
# 5. RENDER RESULTS & CHATBOT (From Memory)
# ==========================================
if st.session_state.run_complete:
    st.markdown("---")
    st.subheader("📈 Gene Expression Volcano Plot")
    st.plotly_chart(st.session_state.volcano_fig, use_container_width=True, key="bottom_volcano_plot")
    
    with st.expander("🔍 View the AI's Strategic Plan"):
        for step in st.session_state.plan:
            st.write(f"- {step}")
            
    # --- NEW: PATHWAY VISUALIZATION ---
    pathway_info = st.session_state.get("pathway_data", {})
    if isinstance(pathway_info, dict) and pathway_info.get("status") == "Success":
        st.markdown("### 🕸️ Enriched Biological Pathways (KEGG)")
        pathways = pathway_info.get("pathways", [])
        
        if pathways:
            # Convert to DataFrame for Plotly
            pw_df = pd.DataFrame(pathways)
            # -log10 transform the p-value for better visualization
            pw_df['Significance Score (-log10 p-value)'] = -np.log10(pw_df['p_value'] + 1e-10)
            
            # Draw horizontal bar chart
            # Draw horizontal bar chart
            pw_fig = px.bar(
                pw_df, 
                x='Significance Score (-log10 p-value)', 
                y='pathway', 
                orientation='h',
                title="Top Associated KEGG Pathways",
                text='overlapping_genes',
                color='Significance Score (-log10 p-value)',
                color_continuous_scale='Sunsetdark' # A warmer, less clunky dark mode palette
            )
            # Clean up the layout, text position, and hover decimals
            pw_fig.update_layout(yaxis={'categoryorder':'total ascending'}, height=300, margin=dict(l=20, r=20, t=40, b=20))
            pw_fig.update_traces(
                textposition='inside', 
                textfont=dict(color='white'),
                hovertemplate="<b>%{y}</b><br>Score: %{x:.2f}<br>Genes: %{text}<extra></extra>"
            )
            st.plotly_chart(pw_fig, use_container_width=True)
            # --- NEW: GSEA MOUNTAIN PLOTS ---
            gsea_obj = st.session_state.get("gsea_obj")
            if gsea_obj:
                st.markdown("### 🏔️ GSEA Enrichment Signatures")
                st.info("These 'Mountain Plots' visualize how the math engine detected the biological shift. The barcode lines represent where the pathway genes fall across the entire tumor genome. A peak on the left means the pathway is strongly upregulated; a trough on the right means it is suppressed.")
                
                # Create columns for up to 3 plots side-by-side
                plot_cols = st.columns(min(3, len(pathways)))
                
                for i, pw in enumerate(pathways[:3]):
                    term = pw["pathway"]
                    try:
                        res_dict = gsea_obj.results[term]
                        res_array = res_dict.get('RES') if 'RES' in res_dict else res_dict.get('res')
                        
                        axes = gseaplot(
                            rank_metric=gsea_obj.ranking, 
                            term=term,
                            hits=res_dict['hits'],
                            nes=res_dict['nes'],
                            pval=res_dict['pval'],
                            fdr=res_dict['fdr'],
                            RES=res_array
                        )
                        
                        if isinstance(axes, list):
                            fig = axes[0].figure
                        elif hasattr(axes, 'figure'):
                            fig = axes.figure
                        else:
                            fig = plt.gcf()
                            
                        with plot_cols[i]:
                            st.pyplot(fig)
                            plt.close(fig) # <-- NEW: Clears the canvas so the next plot is pristine!
                            
                    except Exception as e:
                        with plot_cols[i]:
                            st.error(f"Failed to plot: {term}\nError: {str(e)}")
        else:
            st.info("No statistically significant pathways found for these targets.")

    # --- NEW: TUMOR BOARD TRANSCRIPT ---
    consensus = st.session_state.agent_state.get("expert_consensus", "")
    if consensus:
        with st.expander("🧑‍⚕️ View Raw Tumor Board Debate (Pathologist vs. Oncologist)"):
            st.info("This is the internal reasoning generated by the multi-agent experts before the Medical Writer synthesized the final report.")
            st.markdown(consensus)

    if st.session_state.get("baseline_report"):
        st.markdown("### ⚖️ Head-to-Head Comparison")
        tab1, tab2 = st.tabs(["🤖 OmicsGPT Agent (RAG + Tools)", "⚠️ Vanilla LLM Baseline (No Tools)"])
        
        with tab1:
            st.info("✅ This report was autonomously written by the Medical Writer LLM based solely on validated OpenTargets, OncoKB, and PubMed RAG data.")
            st.markdown(st.session_state.final_report)
            
        with tab2:
            st.warning("🚨 **CAUTION:** This is a standard 'Vanilla' LLM output. It cannot browse the internet, query APIs, or read your custom PDF protocols. It is highly prone to hallucinating clinical trials, mechanisms of action, and outdated drug approvals. It is provided for baseline comparison only.")
            st.markdown(st.session_state.baseline_report)
    else:
        st.markdown("### 📄 Final Synthesized Clinical Report")
        st.info("✅ This report was autonomously written by the Medical Writer LLM based solely on validated tool data.")
        st.markdown(st.session_state.final_report)
    
    # --- NEW: THE CLINICAL AUDIT TRAIL & BIBLIOGRAPHY ---
    st.markdown("### 📚 Reference Library & Evidence Audit")
    
    # 0. THE GLASS BOX PROVENANCE
    used_evidence = st.session_state.agent_state.get("gathered_evidence", [])
    if used_evidence:
        with st.expander("🔍 View AI Semantic Search Algorithm (Provenance)"):
            st.info("Unlike traditional black-box AI search engines, this pipeline uses a deterministic 'Glass Box' methodology combining broad E-Utilities retrieval with FAISS semantic embedding.")
            for g_data in used_evidence:
                provenance = g_data.get("evidence", {}).get("PubMed_Provenance", [])
                if provenance:
                    st.markdown(f"#### **Search Strategy for {g_data['gene']}**")
                    for step in provenance:
                        st.markdown(f"- {step}")
    
    # 1. Show the Papers that WERE used
    used_evidence = st.session_state.agent_state.get("gathered_evidence", [])
    has_kept_papers = False
    
    with st.expander("✅ PubMed Literature Included in Synthesis"):
        for g_data in used_evidence:
            papers = g_data.get("evidence", {}).get("PubMed", {}).get("papers", [])
            if papers:
                has_kept_papers = True
                st.markdown(f"**Target: {g_data['gene']}**")
                for p in papers:
                    st.markdown(f"- **PMID {p['PMID']}**: *{p['Title']}*")
        
        # NEW: If the human or AI threw everything in the trash, print this message!
        if not has_kept_papers:
            st.info("No experimental literature passed the AI quality filter for inclusion. The report relies entirely on systems biology networks and pathway data.")
    
    # 2. Show the Papers that the Human threw out
    discarded = st.session_state.agent_state.get("discarded_evidence", [])
    if discarded:
        with st.expander("🗑️ Manually Filtered (Discarded) Evidence"):
            st.warning("The following literature was manually excluded by the user and hidden from the AI:")
            for idx, paper in enumerate(discarded):
                st.markdown(f"- **{paper['Gene']}** (PMID {paper['PMID']}): *{paper['Title']}*")
                
    # 3. Show the Papers that the AI threw out
    ai_discarded = st.session_state.agent_state.get("ai_filtered_evidence", [])
    if ai_discarded:
        with st.expander("🤖 AI Pre-Filtered Literature (Auto-Discarded)"):
            st.info("The AI evaluated up to 10 papers per gene. The following papers scored < 5 and were automatically excluded.")
            for doc in ai_discarded:
                st.markdown(f"- **{doc['Gene']}** (Score: {doc['Score']}): *{doc['Title']}* - Reason: `{doc['Reason']}`")
    
    # --- EXPORT MENU (HTML & DOCX) ---
    st.markdown("### 💾 Export Options")
    
    html_content = markdown.markdown(st.session_state.final_report, extensions=['tables'])
    
    styled_html = f"""
    <html>
        <head>
            <style>
                body {{ font-family: Arial, sans-serif; line-height: 1.6; color: #333; max-width: 800px; margin: 40px auto; padding: 20px; }}
                h1, h2, h3 {{ color: #2c3e50; border-bottom: 1px solid #eee; padding-bottom: 10px; }}
                a {{ color: #3498db; text-decoration: none; }}
                a:hover {{ text-decoration: underline; }}
                ul {{ margin-bottom: 20px; }}
                li {{ margin-bottom: 8px; }}
            </style>
        </head>
        <body>
            <h1>Clinical AI Orchestrator Report</h1>
            <p><strong>Disease Target:</strong> {cancer_type}</p>
            <hr>
            {html_content}
        </body>
    </html>
    """
    
    doc = Document()
    doc.add_heading(f'Clinical AI Orchestrator Report - {cancer_type}', level=1)
    
    parser = HtmlToDocx()
    parser.add_html_to_document(html_content, doc)
    
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0) 
    
    col_down1, col_down2 = st.columns(2)
    
    with col_down1:
        st.download_button(
            label="🌐 Download as HTML (Browser/PDF)",
            data=styled_html,
            file_name=f"{cancer_type}_Clinical_Report.html",
            mime="text/html",
            use_container_width=True
        )
        
    with col_down2:
        st.download_button(
            label="📄 Download as Word Document (.docx)",
            data=doc_buffer,
            file_name=f"{cancer_type}_Clinical_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

# --- MULTI-MODAL VISUAL ANALYTICS DASHBOARD ---
    st.markdown("---")
    st.subheader("📊 Multi-Modal Target Analytics")
    st.info("Select a target to simultaneously visualize its macroscopic biological neighborhood (Network) and its microscopic physical vulnerabilities (3D Structure).")
    
    if st.session_state.get("ai_targets"):
        # 1. Global Target Selector
        viz_target = st.selectbox("🎯 Select Target to Analyze:", st.session_state.ai_targets, key="viz_target_selector")
        analyze_btn = st.button("Generate Dual Visualization", type="primary", use_container_width=True, key="viz_analyze_btn")
        
        if analyze_btn:
            # 2. Side-by-Side Layout
            col_net, col_struct = st.columns(2)
            
            with col_net:
                st.markdown(f"#### 🕸️ Network Hub: {viz_target}")
                with st.spinner("Building physics network..."):
                    edges = fetch_visual_network(viz_target, max_nodes=15)
                    if edges:
                        net = build_pyvis_graph(viz_target, edges)
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".html") as tmp_file:
                            net.save_graph(tmp_file.name)
                            tmp_file_path = tmp_file.name
                        with open(tmp_file_path, 'r', encoding='utf-8') as HtmlFile:
                            components.html(HtmlFile.read(), height=500)
                        os.remove(tmp_file_path)
                    else:
                        st.warning("Insufficient interaction data in STRING DB.")
            
            with col_struct:
                st.markdown(f"#### 🧬 Structural Pockets: {viz_target}")
                with st.spinner("Mapping 3D coordinates..."):
                    selected_gene_data = next((g for g in st.session_state.agent_state.get("significant_genes", []) if g["hugo"] == viz_target), None)
                    mutation_str = selected_gene_data.get("alteration") if selected_gene_data and "DNA" in selected_gene_data.get("source", "") else None

                    uniprot_id = get_uniprot_id(viz_target)
                    if uniprot_id:
                        struct_string, struct_format = fetch_alphafold_structure(uniprot_id)
                        if struct_string:
                            residues_to_highlight = None
                            if mutation_str:
                                residues_to_highlight = extract_residue_number(mutation_str)
                                st.caption(f"🔴 Highlighting DNA Mutation: **{mutation_str}**")
                            else:
                                residues_to_highlight = get_uniprot_binding_sites(uniprot_id)
                                if residues_to_highlight:
                                    st.caption(f"🔴 Autonomously highlighted **{len(residues_to_highlight)} Active Site residues**.")
                                else:
                                    st.caption("🌈 No pockets found. Displaying structural confidence map.")
                            
                            viewer = render_mutated_protein(struct_string, file_format=struct_format, highlight_residues=residues_to_highlight)
                            showmol(viewer, height=500, width=800)
                        else:
                            st.warning("Failed to download AlphaFold structure.")
                    else:
                        st.warning("Could not map to UniProt ID.")
    else:
        st.warning("No targets available to visualize.")

# --- BENCH-TO-CLOUD VALIDATION DESIGNER ---
    st.markdown("---")
    st.subheader("🧪 Bench-to-Cloud Validation Designer")
    st.info("Translate your in-silico findings into immediate wet-lab action. Generate hypothetical CRISPR knockout sgRNAs and qPCR primers for your selected targets.")
    
    if st.session_state.get("ai_targets"):
        b2c_col1, b2c_col2 = st.columns([1, 3])
        
        with b2c_col1:
            lab_target = st.selectbox("🧬 Select Target for Wet-Lab:", st.session_state.ai_targets, key="lab_select")
            design_btn = st.button("Generate Lab Manifest", type="primary", use_container_width=True)
            
        with b2c_col2:
            if design_btn:
                with st.spinner(f"🤖 AI Molecular Biologist designing experiment for {lab_target}..."):
                    manifest = design_validation_experiment(lab_target, cancer_type)
                    
                if manifest:
                    st.success("✅ Experimental Manifest Generated!")
                    st.markdown(f"**Rationale:** {manifest['rationale']}")
                    
                    st.markdown("#### ✂️ CRISPR-Cas9 sgRNA Designs")
                    st.warning("⚠️ **Clinical Disclaimer:** These sequences are AI-generated for structural planning. You MUST verify them against the human reference genome using Benchling or IDT before ordering.")
                    sgrna_df = pd.DataFrame(manifest['sgrnas'])
                    st.dataframe(sgrna_df, use_container_width=True, hide_index=True)
                    
                    st.markdown("#### 🧬 qPCR Validation Primers")
                    primer_df = pd.DataFrame(manifest['primers'])
                    st.dataframe(primer_df, use_container_width=True, hide_index=True)
                    
                    # --- THE EXPORT IMPROVEMENT ---
                    st.markdown("#### 📥 Export to Vendor")
                    # Combine DataFrames for a single CSV export
                    sgrna_export = sgrna_df.rename(columns={"target_exon": "Name", "sequence": "Sequence"})
                    sgrna_export["Type"] = "sgRNA"
                    primer_export = primer_df.rename(columns={"target": "Name", "forward": "Sequence"})
                    primer_export["Type"] = "Forward Primer"
                    primer_rev_export = primer_df.rename(columns={"target": "Name", "reverse": "Sequence"})
                    primer_rev_export["Type"] = "Reverse Primer"
                    
                    export_df = pd.concat([sgrna_export[["Name", "Type", "Sequence"]], primer_export[["Name", "Type", "Sequence"]], primer_rev_export[["Name", "Type", "Sequence"]]])
                    
                    st.download_button(
                        label="📄 Download IDT/GenScript Plate Manifest (.CSV)",
                        data=export_df.to_csv(index=False).encode('utf-8'),
                        file_name=f"{lab_target}_WetLab_Manifest.csv",
                        mime="text/csv",
                        use_container_width=True
                    )
                else:
                    st.error("AI failed to generate a valid manifest. Please try again.")
            else:
                st.info("👈 Select a target and click 'Generate Lab Manifest'.")

    # --- INTERACTIVE CHATBOT ---
    st.markdown("---")
    st.subheader("💬 Discuss the Findings")
    st.write("Ask follow-up questions about the clinical trials, specific drugs, or resistance mechanisms mentioned above.")
    
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
            
    if prompt := st.chat_input("E.g., What is the mechanism of action for CL-387785?"):
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)
            
        with st.chat_message("assistant"):
            with st.spinner("Thinking..."):
                chat_llm = ChatOpenAI(model="gpt-5.2", temperature=0.2, api_key=openai_key)
                
                chat_sys_msg = f"You are a helpful oncology assistant. Answer the user's questions strictly based on the following report:\n\n{st.session_state.final_report}"
                
                messages = [SystemMessage(content=chat_sys_msg)]
                for m in st.session_state.messages:
                    if m["role"] == "user": messages.append(HumanMessage(content=m["content"]))
                    else: messages.append(AIMessage(content=m["content"]))
                    
                response = chat_llm.invoke(messages)
                st.markdown(response.content)
                
        st.session_state.messages.append({"role": "assistant", "content": response.content})

# --- SIDEBAR: SYSTEM TELEMETRY & DEVELOPER TOOLS ---
st.sidebar.markdown("---")
st.sidebar.markdown("### ⚙️ System Telemetry")
st.sidebar.metric(label="Total Tokens Used", value=f"{st.session_state.total_tokens:,}")

# Track Active Agents based on the LangGraph orchestration
st.sidebar.metric(label="Active Agents", value="4 (Multi-Agent System)")

# Explicit RAG Tracking for Transparency
try:
    rag_status = "Active (FAISS/PubMed/PDF)" if uploaded_pdf is not None else "Active (FAISS/PubMed)"
except NameError:
    rag_status = "Active (FAISS/PubMed)"
st.sidebar.write(f"**RAG Pipeline:** {rag_status}")

st.sidebar.markdown("---")
st.sidebar.markdown("### 🛠️ Developer Tools")
st.sidebar.checkbox("Enable Mock Mode (Bypass APIs & Compute)", value=False, key="use_mock_mode", help="Uses deterministic dummy data for rapid UI styling.")